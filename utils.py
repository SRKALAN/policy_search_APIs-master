# -*- coding: cp1252 -*-
from docx.enum.text import WD_COLOR_INDEX
from itertools import combinations
from docx import Document
import statistics
import time
import numpy as np
import base64
import fitz
import sys
import re
from nltk.corpus import stopwords
from sklearn.cluster import MeanShift, estimate_bandwidth
from sklearn.feature_extraction.text import CountVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from nltk import sent_tokenize

from azure.storage.blob import BlobServiceClient
from io import BytesIO
from rapidfuzz import fuzz
from rapidfuzz import process
from load import init
import time
from nltk import word_tokenize
import string
from string import punctuation


stop = set(stopwords.words('english') + list(string.punctuation))
vectorizer = CountVectorizer()


env=init()
blob_connection_string=env["blob_connection_string"]
container_name_wordings=env["container_name_wordings"]
container_name_documents=env["container_name_documents"]
blob_service_client = BlobServiceClient.from_connection_string(blob_connection_string)
container_client_wordings = blob_service_client.get_container_client(container_name_wordings)
container_client_documents = blob_service_client.get_container_client(container_name_documents)

stop_words = set(stopwords.words('english'))
split_line_re          =re.compile(r'\(.\)|\s.\)|\)|;|:|\.')



rem_lst=[")","(","AND","OR"]
clause_list=['AND','OR']
# Function which returns subset or r length from n

def rSubset(arr, r):
    return list(combinations(arr, r))



class Utils:

    @staticmethod
    def normalise(x,xmin,xmax):
        return (x/xmax)*100


    @staticmethod
    def split_lines(text):
        lines=split_line_re.split(text)
        return lines

    @staticmethod
    def search_text(text,search,num=3):
    #print(num)
        search = re.sub('[.;,":-]$', '', search)
        search_split=search.split()
        #print(search_split)
        first_words=" ".join(search_split[:num])
        last_word= " ".join(search_split[-num:])

        # croped_text=""
        start_index=text.find(first_words)

        end_index=text.rfind(last_word)

        if start_index!=-1 and end_index!=-1:

            if end_index<start_index:
                return start_index,len(text)
            else:
                return start_index,end_index+len(last_word)
        elif end_index!=-1 :

            return 0,end_index+len(last_word)
        elif start_index!=-1 :

            return start_index,len(text)
        if num>2:
            return Utils.search_text(text,search,num-1)
        else:

            return 0,len(text)




    @staticmethod
    def clause_search(text,search):
        start_time = time.time()

        text =re.sub("\s{2,}"," ",text)
        text=re.sub("\n",'',text)

        if text.find(search)!=-1:
            score=100
            highlight_markup="<em>{}</em>".format(search)
            # index=text.find(search)
            # display_text = text[index-20:index+len(search)+20]
            highlighted_text = text.replace(search,highlight_markup)
            # print(search,'......',score,' exact\n\n\n')
            return highlighted_text,score

        elif (fuzz.partial_ratio(search.lower(),text.lower())>=90):
            

            fsearch=search.replace('"','').replace('(','').replace(')','').replace(';',' ').replace(':',' ').replace('.',' ').lower()
            ftext=text.lower().replace('-','')

            if (fsearch.split()[0] in  ftext) and (fsearch.split()[-1] in  ftext):

                flist_Search=[ ftext[y:z].lower()+fsearch.split()[-1].lower() for y in [m.start() for m in re.finditer(fsearch.split()[0], ftext)] for z in [m.start() for m in re.finditer(fsearch.split()[-1], ftext)] if ((y<z) and (len(fsearch)*0.7<= z-y)  and (len(fsearch)*1.5>= z-y) and len(set(ftext[y:z].lower().split()+[fsearch.split()[-1]])-set(fsearch.split()) )<=len(set(fsearch.split()))*0.5)]
 
                if flist_Search:
                    ret_phrase,score,_ = process.extractOne(fsearch,flist_Search)


                    if ( len(ret_phrase.split())>=len(fsearch.split())*0.8) and ret_phrase in ftext:
                        
                        f_ind=ftext.find(ret_phrase)
                        hyphon_count_precd=text[:f_ind].count('-')
                        hyphon_count_tot=text[:f_ind+len(ret_phrase)].count('-')
                        ret_phrase =text[f_ind+hyphon_count_precd:f_ind+len(ret_phrase)+hyphon_count_tot]
                        highlight_markup="<em>{}</em>".format(ret_phrase)
                        highlighted_text = text.replace(ret_phrase,highlight_markup)

                        score =Utils.fuzzy_score(search,ret_phrase)
                        # print(search,'.....',ret_phrase,'.....',score,'fuzzy\n\n\n')
                        return highlighted_text,score

        search=search.replace('"','').replace('(','').replace(')','').replace(';',' ').replace(':',' ')
        search_words=search.split()
        search_words=[w for w in search_words if len(w)>2 and w not in stop_words]

        quantile=len(search)*2/len(text)
        if quantile>=1:
            quantile=0.9
        if quantile<0.2:
            quantile=0.2




        first_word,last_word=search_words[0],search_words[-1]

        position_list=[]
        remp_n=[]
        for word in search_words:
            matches = re.finditer(re.escape(word), text)
            matches_positions = [(match.start()) for match in matches]
            position_list.extend(matches_positions)
            remp_n.append([matches_positions,word])
        position_list=sorted(position_list)



        # if bandwidth > 0
        # try:
        #     X = np.array(list(zip(position_list,np.zeros(len(position_list)))), dtype=np.int)
        #     if len(X)<1000:
        #         bandwidth = estimate_bandwidth(X, quantile=quantile)

        #     else:
        #         print("custom function")
        #         bandwidth = len(X)*0.77
        #     ms = MeanShift(bandwidth=bandwidth, bin_seeding=True)
        #     ms.fit(X)
        #     labels = ms.labels_

        #     labels_unique = np.unique(labels)
        #     n_clusters_ = len(labels_unique)

        #     my_members = labels == 0
        #     cluster_positions=X[my_members, 0]
        #     start_position_adjustment=0.5
        #     end_position_adjustment=1.5
        #     start_clus=min(cluster_positions)
        #     end_clus=max(cluster_positions)

        # except Exception as err:
        #     print(err)

        start_position_adjustment=0.8
        end_position_adjustment=1.2

        cluster_positions=position_list

        cluster_master=[]
        cluster=[]
        for indx, pos in enumerate(position_list):
            if cluster:
                if pos-cluster[-1] <80:
                    cluster.append(pos)
                else:
                    cluster_master.append(cluster)
                    cluster=[]

            else:
                cluster.append(pos)
        if cluster_master:
            cluster_positions=max(cluster_master, key = len)

        else:
            return False,0





        if len(cluster_positions)<2:
            return False,0


        #
        start_clus=min(cluster_positions)
        end_clus=max(cluster_positions)
        for zx in remp_n:
            if end_clus in zx[0]:
                end_adj=len(zx[1])
                end_word=zx[1]
                break
        cluster_text=text[start_clus:end_clus+end_adj]







        start_display=text[:int(round(start_position_adjustment*start_clus))].rfind(' ')
        end_display=text[int(round(end_position_adjustment*end_clus)):].find(' ')+int(round(end_position_adjustment*end_clus))
        if start_display==-1:
            start_display=0
        if end_display==-1:
            end_display=len(text)
        display_text=text[start_display:end_display]
        if len(display_text)/len(cluster_text) >2.5 :

            display_text=cluster_text



        # print(display_text,'\n')

        # print("cluster_text",cluster_text)
        s,e= Utils.search_text(cluster_text.lower(),search.lower(),num=3)





        if s==0 and e ==len(cluster_text):


            s=cluster_text.lower().find(first_word.lower())
            e=cluster_text.lower().rfind(last_word.lower())

            s_index=1
            while (e ==-1 or 'means' in first_word  ) and s_index <3  :
                first_word=search_words[s_index]
                s=cluster_text.find(first_word)


                s_index+=1
            e_index=1
            while (e ==-1 or 'means' in last_word   )and  e_index <4  :
                last_word=search_words[-(e_index)]
                e=cluster_text.rfind(last_word)

                e_index+=1
            e+=len(last_word)
        # print(s,e)

        if e<s:
            s=None
            e=None
        if s==-1:
            s=None
        if e==-1:
            e=None

        #if e:
        #    e


        clause=cluster_text[s:e]




        if len(cluster_text)/len(clause) >6.5 :
            clause=cluster_text






        score=Utils.fuzzy_score(search,clause)


        # print(search,'....',clause,'.....',score,'meanshift\n\n')
        






        if score<50:

            return False,0
        highlight_markup="<em>{}</em>".format(clause)

        highlighted_text=display_text.replace(clause,highlight_markup)






        return highlighted_text,score


    @staticmethod
    def count_cosine(vectorizer,text1,text2):
        tot_text=[text1,text2]
        vectorizer.fit(tot_text)
        v1=vectorizer.transform([text1])
        v2=vectorizer.transform([text2])

        score=cosine_similarity(v1,v2)[0][0]
        return score
    @staticmethod
    def fuzzy_score(text1,text2):
        text2=re.sub('[^A-Za-z0-9]+', ' ', text2)
        text1=re.sub('[^A-Za-z0-9]+', ' ', text1)
        score=fuzz.ratio(text1.lower(),text2.lower())
        return score

    @staticmethod
    def highlight_def1(text,phrases):
        flag=False
#     pos_start=text.find('###')
        pos_def=text.find('@@')
        pos_end=text.find('%%%')
        if not pos_end:
            pos_end=text.find('.....')
        term=text[:pos_def]
        definition=text[pos_def:pos_end]
    #     definition
        text1=term+text[pos_end:]
        for phrase in phrases:
            text1=text1.replace(phrase,'<em>'+phrase+'</em>')
            if( phrase.lower() in definition.lower()) and phrase.lower()!=term.lower():
                flag=True
        if flag:
            text1=text1.replace(term,'<em>'+term+'</em>')
        text1=text1.replace('###','')
        text1=text1.replace('%%%','')
        return(text1)

    @staticmethod
    def delete_params(result):
        params=['bold_phrases','highlight_words_def','definitions_in_page','definition_text',
        'proximity','custom_score','is_clause','proximity_crop','definition_confidence','exclusion_confidence',
        'condition_confidence','text','azure_score','definitions','highlight_words']
        for param in params:
            del result[param]

        return result







    @staticmethod
    def spell_check_lite(search_string):
        search_string_clean=search_string
        for item in rem_lst:
            search_string_clean=search_string_clean.replace(item,' ')
        search_words=search_string_clean.split()
        if len(search_words)==1:
            search_string=search_words[0]+"~"
            return search_string
        else:
            return search_string






    @staticmethod
    def spell_check(search_string):

        search_string_clean=search_string

        for item in rem_lst:
            search_string_clean=search_string_clean.replace(item,' ')

        search_words=search_string_clean.split()
        search_words_dict={word:word+'~' for word in search_words }
        for word in search_words:
            search_string=search_string.replace(word,search_words_dict[word])
        return search_string

    @staticmethod
    def kephrase_process(phrases):
        phrases=[phrase.replace('  ',' ') for phrase in phrases]
        phrases=list(set(phrases))
        phrases=[phrase for phrase in phrases if len(phrase)>3]
        return phrases



    @staticmethod
    def get_phrase_proximity(highlight_words,phrase,text,super_list):
        return 0


        #making sure the phrases in between words in AND clauses are having maximum proximity
        for master_list in super_list:
            _,cropped_string=Utils.proximity(text,master_list)
            if phrase in cropped_string:
                return 0






        phrase_position=text.find(phrase)
        position_list=[]
        for word in highlight_words:
            matches = re.finditer(re.escape(word), text)
            matches_positions = [match.start() for match in matches]
            position_list.extend(matches_positions)

        if not position_list:
            return 1000

        text_window_lengths=[]
        for position in position_list:
            if phrase_position>position:
                text_window=text[position:phrase_position]
            else:
                text_window=text[phrase_position:position]
            text_window=text_window.split()

            text_window_lengths.append(len(text_window))
        if not text_window_lengths:
            return 1000
        return min(text_window_lengths)



        min_distance=min(distance_list)
        distance_score=1-(min_distance/len(text))
        return distance_score


    @staticmethod
    def word_highlighter(source,phrases):
        word_path="blobs/{}.docx".format(source)
        out_path="out/{}.docx".format(source)
        doc = Document(word_path)
        all_text=''
        for para in doc.paragraphs :
            for phrase in phrases:
                start = para.text.lower().find(phrase.lower())
                if start > -1 :
                    pre = para.text[:start]
                    post = para.text[start+len(phrase):]
                    para.text = pre
                    para.add_run(phrase)
                    para.runs[1].font.highlight_color = WD_COLOR_INDEX.YELLOW
                    para.add_run(post)
        doc.save(out_path)
        with open(out_path, "rb") as word_file:
            encoded_string = base64.b64encode(word_file.read())
        return encoded_string

    @staticmethod
    def distance_group(index_list,distance=2):
        index_list.sort()

        group_index=[[index_list[0]]]
        for i in index_list[1:]:

            if abs(i-group_index[-1][-1])<=distance:

                group_index[-1].append(i)
            else:
                group_index.append([i])
        return group_index

    @staticmethod
    def get_highlightWords(paragraph,phrase_list):

        split_list= [  [i.replace("'s'","") for i in a.split() if i not in stop ]    for a in phrase_list ]
        #split_list= [  [i for i in a.split() if i not in stop ]    for a in phrase_list ]
        highlight_words=[]
        for  words in split_list:

            if len(words)==1:
                highlight_words.extend(words)
                continue


            for line in split_line_re.split(paragraph):

                if all(word.lower() in line.lower()  for word in words  ):

                    sentence_split = line.split()

                    index_list=[]

                    for word in words:

                        matched_words=re.findall(r"{}[^\s]{}".format(word,"{0,2}"), line, flags=re.IGNORECASE)

                        index_list.extend( [i for i, x in enumerate(sentence_split) for word in matched_words if   word.lower() in x.lower()])


                    if index_list:
                        group_index = Utils.distance_group(index_list,3)

                        for word_index in group_index:

                            search_word=" ".join(sentence_split[word_index[0]:word_index[-1]+1] )


                            if all(word.lower()  in search_word.lower() for word in words):
                                highlight_words.append(search_word)

        return highlight_words






    @staticmethod
    def pdf_highlighter(file,phrases,phrases_def,folder,user_id,country_code,type_of_doc,phrase_list,page_num):
        pdf_path="{}/{}/{}.pdf".format(folder,type_of_doc.title(),file)
        blob_client = container_client_wordings.get_blob_client(pdf_path)
        out_path="out/{}_{}.pdf".format(file,user_id)
        print(pdf_path)
        stream_buffer = BytesIO()
        with stream_buffer as download_stream:
            blob_client.download_blob(max_concurrency=3).readinto(download_stream) # just download PDFs from Azure as in-memory streams
            doc = fitz.open(stream=stream_buffer, filetype="pdf")
            print('Streaming...')


        #for page in doc:
        page = doc.loadPage(int(page_num)-1)
        paragraph=page.getText("text")

        if phrases_def:
            for match_text in phrases_def:
                text_instances = page.searchFor(match_text)
                if text_instances:
                    for inst in text_instances:
                        highlight = page.addHighlightAnnot(inst)
                        highlight.setColors({"stroke":(0, 1, 1), "fill":(0.7, 0.7, 0.7)})
                        highlight.update()

        high_light_words = Utils.get_highlightWords(paragraph,phrase_list)
        # print("high_light_words",high_light_words)
        for text in high_light_words:
            text_instances = page.searchFor(text)
            if page.searchFor(text.capitalize() ):
                text_instances.extend( page.searchFor(text.capitalize() ))
            if page.searchFor(text.upper() ):
                text_instances.extend( page.searchFor(text.upper() ))
            if page.searchFor(text.lower() ):
                text_instances.extend( page.searchFor(text.lower() ))
            #print('printing pdf highlightuing strings identified',text_instances,text)
            if text_instances:
                for inst in list(set(text_instances)):
                    highlight = page.addHighlightAnnot(inst)

        for match_text in phrases:

            if len(match_text.split()) <2:
                pass
            else:
                lines=split_line_re.split(match_text)
                for line in lines:
                    text_instances = page.searchFor(line)
                    if text_instances:
                        for inst in text_instances:
                            highlight = page.addHighlightAnnot(inst)
                if not text_instances:
                    words = match_text.split()
                    grouped_words = [' '.join(words[i: i + 4]) for i in range(0, len(words), 4)]
                    for line in grouped_words:
                        text_instances = page.searchFor(line)
                        if text_instances:
                            for inst in text_instances:
                                highlight = page.addHighlightAnnot(inst)





        doc.save(out_path, garbage=4, deflate=True, clean=True)
        with open(out_path, "rb") as pdf_file:
            encoded_string = base64.b64encode(pdf_file.read())
        return encoded_string

    @staticmethod
    def click_pdf_highlighter(source,phrase,user_id):
        pdf_path="out/{}_{}.pdf".format(source,user_id)
        out_path="click_tag/{}_{}.pdf".format(source,user_id)
        doc = fitz.open(pdf_path)
        for page in doc:
            text_instances = page.searchFor(phrase)
            for inst in text_instances:
                highlight = page.addHighlightAnnot(inst)
                highlight.setColors({"stroke":(0, 1, 1), "fill":(0.55, 0.7, 0.55)})
                highlight.update()

        doc.save(out_path, garbage=4, deflate=True, clean=True)
        with open(out_path, "rb") as pdf_file:
            encoded_string = base64.b64encode(pdf_file.read())
        return encoded_string
    @staticmethod
    def click_docx_highlighter(source,phrase):
        word_path="out/{}.docx".format(source)
        out_path="click_tag/{}.docx".format(source)
        doc = Document(word_path)
        all_text=''
        for para in doc.paragraphs :
            start = para.text.lower().find(phrase.lower())
            if start > -1 :
                pre = para.text[:start]
                post = para.text[start+len(phrase):]
                para.text = pre
                para.add_run(phrase)
                para.runs[1].font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                para.add_run(post)
        doc.save(out_path)
        with open(out_path, "rb") as word_file:
            encoded_string = base64.b64encode(word_file.read())
        return encoded_string

    @staticmethod
    def confidence_bow(search_words,excl_list,text):
        page_char_length=len(text)
        position_list=[]
        result=[]
        for word in search_words:
            for match in re.finditer(re.escape(word), text.lower()):
                position_list.append(match.start())
        if excl_list and position_list:
            for pos_ex in excl_list:
                for pos_search in position_list:
                    result.append(abs(pos_search-pos_ex))
            # print(result,'RES LIST')
        try:
            min_distance=min(result)
            # print(min_distance,'MIN')
            if min_distance < 100:
                score=(1-(min_distance/100))*100
                # print(score,'SCORE!!!''\n\n')
            else:
                # print('ELSE BLOCK\n\n')
                score=0
        except Exception as err:
            # print(str(err),'EXCEPT BLOCK\n\n')
            score=0

        return score
    @staticmethod
    def section_confidence_bow(search_words,bow,text):

        lines=split_line_re.split(text.lower())

        section_flag=False
        for line in lines:
            bow_flag=any([word in line for word in bow])
            search_word_flag=any([word in line for word in search_words])
            if bow_flag and search_word_flag:
                section_flag=True
                break



            # print(result,'RES LIST')


        return section_flag

    @staticmethod
    def find_def_in_keys(highlight_words,definitions):
        def_key_list=[i['name'].lower() for i in definitions ]
        highlight_words=[i.lower() for i in highlight_words]

        for def_name in def_key_list:
            for word in highlight_words:
                if len(word) <3 or len(def_name) <3:
                    continue
                if def_name in word or word in def_name:
                    Ratio = fuzz.ratio(def_name.lower(),word.lower())
                    # print(Ratio,'FUZZY MATCH')
                    return True,Ratio
        return False,0


    @staticmethod
    def incremental_count(lob,country,saved_search_container):
        query="SELECT c.inremental_count FROM c WHERE c.lob ='{}' AND  c.country ='{}'".format(lob,country)
        items = list(saved_search_container.query_items(
            query=query,
            enable_cross_partition_query=True
        ))

        try:
            count=max([i['inremental_count'] for i in items])
        except:
            count=0
        count=count+1
        return count



    @staticmethod
    def payload_to_super_list(phrases):
        clauses=[]
        super_list=[]
        master_list=[]
        phrase_list=[]
        for item in phrases:
            block=item['searchInput']
            for phrase in block:
                if len(phrase['input'].split()) >5:
                    clauses.append(phrase['input'])
                    continue

                phrase_list.append(phrase['input'].lower())
                if phrase['checked']:
                    master_list.append(phrase_list)
                    phrase_list=[]
            if phrase_list:
                master_list.append(phrase_list)
                phrase_list=[]
            if not item["checked"]:
                    super_list.append(master_list)
                    master_list=[]
        super_list.append(master_list)
        return super_list,clauses

    @staticmethod
    def payload_to_super_list_n_clauses(phrases):
        clauses=[]
        super_list=[]
        master_list=[]
        phrase_list=[]
        clause_and,clause_or=False,False
        for item in phrases:
            block=item['searchInput']
            for phrase in block:
                if len(phrase['input'].split()) >=6:
                    clauses.append(phrase['input'])
                    if  phrase['checked'] :
                        clause_and =True
                    continue

                phrase_list.append(phrase['input'].lower())
                if phrase['checked']:
                    master_list.append(phrase_list)
                    phrase_list=[]
            if phrase_list:
                master_list.append(phrase_list)
                phrase_list=[]
            if not item["checked"]:
                    super_list.append(master_list)
                    master_list=[]
        super_list.append(master_list)

        return super_list,clauses,clause_and





    @staticmethod
    def findSmallestRange(arr,n,k):
        try:
            ptr = [0 for i in range(501)]
            i, minval, maxval, minrange, minel, maxel, flag, minind = 0,0,0,0,0,0,0,0
            for i in range(k + 1):
                      ptr[i] = 0
            minrange = 10**9
            while(1):
                minind = -1
                minval = 10**9
                maxval = -10**9
                flag = 0

                for i in range(k):

                    if(ptr[i] == n):
                        flag = 1
                        break
                    if(ptr[i] < n and arr[i][ptr[i]] < minval):
                        minind = i # update the index of the list
                        minval = arr[i][ptr[i]]

                    if(ptr[i] < n and arr[i][ptr[i]] > maxval):
                        maxval = arr[i][ptr[i]]


                if(flag):
                    break

                ptr[minind] += 1

                               # updating the minrange
                if((maxval-minval) < minrange):
                    minel = minval
                    maxel = maxval
                    minrange = maxel - minel
        except:
            pass
        return minrange,[(minel,minel+6),(maxel,maxel+9)]


    @staticmethod
    def search_relevance_keyword (words_proximity_count,sub_heading=False):
        match,relevance = "full","High"
        if words_proximity_count < 50:
            pass
        elif words_proximity_count >= 50 and words_proximity_count < 75:
            match,relevance = "partial","High"
        elif words_proximity_count >= 75 and words_proximity_count < 100:
            match,relevance = "partial","Medium"
        else:
            match,relevance = "partial","Low"
        if sub_heading:
            match,relevance = "partial","Low"

        return match,relevance

    @staticmethod
    def search_relevance_clause (score):
        match,relevance = "partial","Low"
        if score>=99:
            match,relevance = "full",""
        elif score >=95:
            match,relevance = "partial","High"
        elif score >=85:
            match,relevance = "partial","Medium"
        return match,relevance



    @staticmethod
    def check_heading_bow(text,word_list,header_pos,section_bow):

        # return True

        matches_positions=[]

        for term in word_list:
            term=term.strip()
            matches = re.finditer(re.escape(term), text)
            matches_positions.extend([match.start() for match in matches])

        # matches =[ re.finditer(re.escape(term.strip()), text) for term in word_list]

        # matches_positions=[match.start() for match in matches]
        # print(word_list)
        # print(matches_positions,header_pos,section_bow)
        for match_pos in matches_positions:
            for section_pos in section_bow:
                if section_pos<match_pos:
                    list_of_nmbrs=set(range(section_pos,match_pos))
                else:
                    list_of_nmbrs=set(range(match_pos,section_pos))


                # print(list_of_nmbrs,header_pos)
                intersection = list_of_nmbrs.intersection(header_pos)

                if  not intersection:
                    return True

        return False
                



    @staticmethod
    def proximity(text,word_list):
        # return 10,'random words'
        '''Tries to find the smallest window in which atleast on term from each of the list is found. '''
        text=re.sub("\s{2,}"," ",text)
        text=text.lower()
        flat_list = [item for sublist in word_list for item in sublist]


        position_list=[]
        split_terms=[]
        for term in flat_list:
            term=term.strip()
            matches = re.finditer(re.escape(term), text)
            matches_positions = [(match.start(),match.start()+len(term)) for match in matches]



            #when two words in the a term is placed different they are considered to be present if they are found with in a window

            if not matches_positions and len(term.split())>1:

                term=term.replace("'s",'')


                words=term.split()
                words=[word for word in words if word not in stop_words]
                required_window_size=len(words)+1

                word_position_list=[]
                word_pos=[]

                for word in words:

                    matches = re.finditer(re.escape(word), text)
                    matches_positions_words_s_e = [(match.start(),match.start()+len(term)) for match in matches]
                    matches_positions_words_start = [i[0] for i in matches_positions_words_s_e]

                    if matches_positions_words_start:
                        word_position_list.append(matches_positions_words_start)
                        word_pos.extend(matches_positions_words_s_e)




                if len(word_position_list)==len(words):
                    try:
                        window,pos_s_e=Utils.findSmallestRange(word_position_list,5,len(word_position_list))
                        cropped_window=text[pos_s_e[0][0]:pos_s_e[-1][1]]

                        window_word_size=len(cropped_window.split())

                    except:

                        window_word_size=500

                    if window_word_size <=required_window_size:

                        split_terms.append(term)
                        position_list.extend(pos_s_e)

            position_list.extend(matches_positions)
        combinations_list=rSubset(position_list,2)



        combinations_list=list(set(combinations_list))
        crop_positions={}

        for start,end in combinations_list:
            if start[0] < end[0] :
                start=start[0]
                end=end[1]

                cropped_string=text[start:end]
            else:
                start=start[1]
                end=end[0]
                cropped_string=text[end:start]
            check_all=[]

            for group in word_list:
                group_presense=[]
                for word in group:
                    word=word.replace("'s",'')

                    flag=False
                    if word in  split_terms:
                        if any(split_word in cropped_string for split_word in word.split() if split_word not in stop_words):
                            flag=True
                    if word in cropped_string:
                        flag=True

                    group_presense.append(flag)


                if any(group_presense):
                    check_all.append(True)
                else:
                    check_all.append(False)

            if all(check_all):
                crop_positions[cropped_string]=len(cropped_string.split())

        if not crop_positions:
            return len(text.split()),''
        cropped_string=min(crop_positions,key=crop_positions.get)
        proximity=crop_positions[cropped_string]


        return proximity,cropped_string



    @staticmethod
    def single_level_proximity(text,word_list):
        text=re.sub("\s{2,}"," ",text)
        text=text.lower()
        flat_list = [item for sublist in word_list for item in sublist]
        # print(text, word_list)
        position_list=[]
        split_terms=[]

        for term in flat_list:
            if term in text:
                return 0,'Not applicable'
        window_dict={}
        for term in flat_list:

            words=term.split()
            words=[word for word in words if word not in stop_words]
            if  all(wrd.lower() in text.lower() for wrd in words):
                word_position_list=[]

                for word in words:
                    matches = re.finditer(re.escape(word), text)
                    matches_positions_words_s_e = [(match.start(),match.start()+len(term)) for match in matches]
                    matches_positions_words_start = [i[0] for i in matches_positions_words_s_e]
                    # print(word,matches_positions_words_start)
                    if matches_positions_words_start:
                        word_position_list.append(matches_positions_words_start)


                if len(word_position_list)>1:
                    try:
                        window,pos_s_e=Utils.findSmallestRange(word_position_list,5,len(word_position_list))
                        if window <30:
                            window_text=text[pos_s_e[0][0]:pos_s_e[1][1]]
                            window_dict[window_text]=window
                    except Exception as err:
                        print(err)
                        pass

        if window_dict:
            window_text = max(window_dict,key=window_dict.get)
            window=window_dict[window_text]

            return window/3,window_text
        else:
            return len(text.split()),None

    @staticmethod
    def proximity_header_excluded(text,word_list,header_pos):
        # return 10,'random words'
        '''Tries to find the smallest window in which atleast on term from each of the list is found. '''
        text=re.sub("\s{2,}"," ",text)
        text=text.lower()
        # print(text)
        flat_list = [item for sublist in word_list for item in sublist]


        position_list=[]
        split_terms=[]
        header_exl=[]
        for term in flat_list:
            term=term.strip()
            matches = re.finditer(re.escape(term), text)
            matches_positions = [(match.start(),match.start()+len(term)) for match in matches]



            #when two words in the a term is placed different they are considered to be present if they are found with in a window

            if  len(term.split())>1:

                term=term.replace("'s",'')


                words=term.split()
                words=[word for word in words if word not in stop_words]
                required_window_size=len(words)+3

                word_position_list=[]
                word_pos=[]

                for word in words:

                    matches = re.finditer(re.escape(word), text)
                    matches_positions_words_s_e = [(match.start(),match.start()+len(term)) for match in matches]
                    matches_positions_words_start = [i[0] for i in matches_positions_words_s_e]
                    # print(word,matches_positions_words_s_e)


                    if matches_positions_words_start:
                        word_position_list.append(matches_positions_words_start)
                        word_pos.extend(matches_positions_words_s_e)




                # position_list.extend(word_pos)
                if len(word_position_list)>=len(words):
                    try:
                        window,pos_s_e=Utils.findSmallestRange(word_position_list,5,len(word_position_list))
                        cropped_window=text[pos_s_e[0][0]:pos_s_e[-1][1]]

                        window_word_size=len(cropped_window.split())
                        # print(window_word_size)   

                    except:

                        window_word_size=500

                    if window_word_size <=required_window_size:

                        split_terms.append(term)
                        position_list.extend(pos_s_e)

            position_list.extend(matches_positions)
        combinations_list=rSubset(position_list,2)

        # print(position_list)





        combinations_list=list(set(combinations_list))
        # print(position_list)
        crop_positions={}

        for start,end in combinations_list:
            if start[0] < end[0]:
                cropped_string=text[start[0]:end[1]]

                if not any(pos >= start[0] and pos <= end[1] for pos in header_pos) or cropped_string.find('.')==-1  :
                    start=start[0]
                    end=end[1]
                    header_flag=False
                else:
                    header_flag=True
            elif end[0]<start[0]:
                cropped_string=text[end[0]:start[1]]
                if not any([pos >= end[0] and pos <= start[1] for pos in header_pos]) or cropped_string.find('.')==-1:
                    start=start[1]
                    end=end[0]
                    header_flag=False
                else:
                    header_flag=True
            else:
                continue

            check_all=[]
            
            for group in word_list:
                group_presense=[]
                for word in group:

                    flag=False
                    if all(split_word in cropped_string for split_word in word.split() if split_word not in stop_words):
                        flag=True
                        # print('first')
                    if word in cropped_string:
                        # print('second',word)
                        flag=True
                        # print(word)
                    group_presense.append(flag)
                    # if flag==False and word=='cancellation of lease' and  'lease' in cropped_string and 'cancellation' in cropped_string  :
                    #     print(split_terms)


                if any(group_presense):
                    check_all.append(True)
                else:
                    check_all.append(False)
            if all(check_all):
                crop_positions[cropped_string,header_flag]=len(cropped_string.split())

        # print('....',crop_positions)



        if not crop_positions:
            return len(text.split()),'',False
        crop_positions_head={i:j for i,j in crop_positions.items() if i[1]==True}
        crop_positions_not_head={i:j for i,j in crop_positions.items() if i[1]==False}


        if crop_positions_not_head:
            cropped_string,header_flag=min(crop_positions_not_head,key=crop_positions_not_head.get)
        else:
            cropped_string,header_flag=min(crop_positions_head,key=crop_positions_head.get)

        proximity=crop_positions[cropped_string,header_flag]
        # print('cropppp....',cropped_string,header_flag)
        if '.' in cropped_string and  'means' in cropped_string :
            header_flag=True


        return proximity,cropped_string,header_flag

    @staticmethod
    def doc_pdf_highlighter(file,lob,phrases,phrases_def,user_id,country,phrase_list,page_num,section_bow):

        # docname,highlight_words,highlight_words_def,user_id,country,phrase_list,page_num
        # if country=='gb' or country=='GB':
        country_code='GB'

        lob=lob[0]
        if lob=="fire":
            lob='Fire'
        #pdf_path="{}-{}/{}.pdf".format(country_code,lob.title(),file) I commented this out as a test this is the original
        pdf_path="{}-{}/{}.pdf".format(country_code,lob,file) # just getting the first value as a test
        # print(pdf_path,'PDF PATH')
        blob_client = container_client_documents.get_blob_client(pdf_path)
        out_path="out/{}_{}.pdf".format(file,user_id)
        stream_buffer = BytesIO()
        with stream_buffer as download_stream:
            blob_client.download_blob(max_concurrency=3).readinto(download_stream) # just download PDFs from Azure as in-memory streams
            doc = fitz.open(stream=stream_buffer, filetype="pdf")
            print('Streaming...')


        #for page in doc:
        page = doc.loadPage(int(page_num)-1)
        paragraph=page.getText("text")
        if section_bow is not None:
            for bow in section_bow:
                text_instances = page.searchFor(bow)

                if text_instances:
                    for inst in text_instances:
                        highlight = page.addHighlightAnnot(inst)
                        highlight.setColors({"stroke":(0, 1, 1), "fill":(0.48, 0.9, 0)})
                        highlight.update()
        if phrases_def:
            for match_text in phrases_def:
                text_instances = page.searchFor(match_text)
                if text_instances:
                    for inst in text_instances:
                        highlight = page.addHighlightAnnot(inst)
                        highlight.setColors({"stroke":(0, 1, 1), "fill":(0.7, 0.7, 0.7)})
                        highlight.update()

        high_light_words = Utils.get_highlightWords(paragraph,phrase_list)
        # print("high_light_words",high_light_words)
        for text in high_light_words:
            text_instances = page.searchFor(text)
            if text_instances:
                for inst in text_instances:
                    highlight = page.addHighlightAnnot(inst)

        for match_text in phrases:

            if len(match_text.split()) <2:
                pass
            else:
                lines=split_line_re.split(match_text)
                for line in lines:

                    if len(line)>5:
                        text_instances = page.searchFor(line)
                        
                        if text_instances:
                            for inst in text_instances:
                                highlight = page.addHighlightAnnot(inst)
                        if not text_instances:
                            words = match_text.split()
                            grouped_words = [' '.join(words[i: i + 5]) for i in range(0, len(words), 5)]

                            small_text_instances=[]
                            smaller_text_instances=[]
                            for line in grouped_words:
                                if len(line)>10:
                                    text_instances=page.searchFor(line)
                                    if text_instances:
                                        small_text_instances.extend(text_instances)
                                    else:
                                        clean_line=(' ').join([i for i in line.split() if len(i)>1])
                                        small_text_instances.extend(page.searchFor(clean_line))


                                # else:
                                #     print(line,'@@@@@@')
                                #     smaller_text_instances.extend(page.searchFor(line))


                            small_text_instances_y=[inst[1] for inst in small_text_instances]
                            # smaller_text_instances_y=[inst[1] for inst in smaller_text_instances]
                            # print(small_text_instances_y)
                            # print(small_text_instances_y)
                            small_text_instances_y=sorted(small_text_instances_y)


                            cluster_master=[]
                            cluster=[]
                            for  pos in small_text_instances_y:
                                if cluster:
                                    if abs(pos-cluster[-1] )<40:
                                        cluster.append(pos)
                                    else:
                                        cluster_master.append(cluster)
                                        cluster=[]

                                else:
                                    cluster.append(pos)
                            # tot_text_instances=small_text_instances_y.extend(smaller_text_instances_y)
                            cluster_master.append(cluster)

                            cluster_max= max(cluster_master, key = lambda i: len(i))
                            text_instances_cluster=[inst for inst in small_text_instances if inst[1] in cluster_max ]
                            for inst in text_instances_cluster:
                                highlight = page.addHighlightAnnot(inst)





        doc.save(out_path, garbage=4, deflate=True, clean=True)
        with open(out_path, "rb") as pdf_file:
            pdf = pdf_file.read()
        return pdf
